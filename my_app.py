from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document ()

# profile picture
document.add_picture(
    'me.png',
    width = Inches(2.0)
)

# name, phone number and email details
name = input('What is your name? ')
speak('Welcome ' + name + ' How are you today? ')
speak('What is your phone number?')
phone_number = input('What is your phone number? ')
email = input('Type your email address ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
about_me = input('Tell me about yourself ')
document.add_paragraph(about_me)

# Work experiences
document.add_heading('Work Experiences')
p = document.add_paragraph()

company = input('Type the company ')
from_date = input('Type the starting date ')
to_date = input('to date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic

experience_details = input(
    'Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Type the company ')
        from_date = input('Type the starting date ')
        to_date = input('to date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic

        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
skill = input('Enter a skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigoscode and Intuit Quickbooks course project"

document.save('cv.docx')
